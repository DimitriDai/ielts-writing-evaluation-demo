
import os
import re
import json
import docx
import requests
from typing import Literal
from docx import Document

# 自动获取 API KEY（兼容 streamlit 和环境变量）
try:
    import streamlit as st
    DEEPSEEK_API_KEY = st.secrets["DEEPSEEK_API_KEY"]
except:
    DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")

def load_text_from_file(filepath: str) -> str:
    if filepath.lower().endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    elif filepath.lower().endswith(".docx"):
        doc = docx.Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    else:
        raise ValueError("仅支持 .txt 或 .docx 文件")

def guess_task_type(text: str) -> Literal["Task 1", "Task 2"]:
    task1_keywords = ["the chart", "the table", "graph shows", "diagram illustrates"]
    if any(kw in text.lower() for kw in task1_keywords):
        return "Task 1"
    return "Task 2"

def get_feedback_from_deepseek(text: str, api_key: str, task_type="Task 2") -> str:
    headers = {
        "Authorization": f"Bearer {api_key or DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }

    full_prompt = f"""你是一位雅思写作考官，请你根据下列规则对学生的作文进行逐句反馈与优化：
1. 逐句指出每一句中是否存在语法、拼写、词汇、搭配、表达不自然等错误。
2. 对有问题的句子，逐句给出修正后的句子，但不改变原句的表达内容。
3. 请根据 I+1 教学原则，对每个段落进行整体优化，生成提升表达的英文段落，整理放在文档最后。如果字数不达标，请给主体段各增加一个例子，每个例子至少2句话。
4. 最后，参考预打分，给雅思作文每一项给出打分（Task Response、Coherence and Cohesion、Lexical Resource、Grammatical Range and Accuracy），并说明每项分数的理由。
5. 所有优化段落统一放在文末【段落优化】部分，并简要说明优化依据。 
6. 输出格式要求如下： 
- 每个句子前加上阿拉伯数字序号；
- 每个句子后依次写出【错误】、【修正】（若无错误，写“无”）；
- 在“段落优化”部分前加上“段落优化：”，中文标点符号，只显示优化后的段落即可。

请按以上格式处理以下学生作文（类型为 {task_type}）：
{text.strip()}
"""
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": full_prompt}],
        "temperature": 0.3
    }

    url = "https://api.deepseek.com/v1/chat/completions"
    response = requests.post(url, headers=headers, json=payload)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def clean_feedback_text(text: str) -> str:
    return text.replace("*", "").replace("#", "").strip()

def save_feedback_to_docx(text: str, output_path: str, task_type: str = None, band_scores: dict = None):
    doc = Document()
    if task_type:
        doc.add_heading(f"作文类型：{task_type}", level=1)

    def clean_feedback_text(t): return t.replace("*", "").replace("#", "").strip()
    text = clean_feedback_text(text)
    doc.add_paragraph(text)

    doc.add_heading("评分建议及说明", level=1)

    dimensions = ["Task Response", "Coherence and Cohesion", "Lexical Resource", "Grammatical Range and Accuracy"]

    doc.add_heading("机器评分：", level=2)
    total = 0
    for dim in dimensions:
        score = band_scores.get(dim, {}).get("score", "-")
        reasons = band_scores.get(dim, {}).get("reasons", [])
        reason_text = "；".join(reasons) if reasons else "无明确降分依据，默认评分"
        doc.add_paragraph(f"{dim}: {score} 分（{reason_text}）")
        if isinstance(score, (int, float)) or (isinstance(score, str) and score.replace('.', '', 1).isdigit()):
            total += float(score)

    doc.add_heading("人工评分（教师填写）", level=2)
    doc.add_paragraph("Task Response: ____")
    doc.add_paragraph("Coherence and Cohesion: ____")
    doc.add_paragraph("Lexical Resource: ____")
    doc.add_paragraph("Grammatical Range and Accuracy: ____")
    doc.add_paragraph("📌 最终得分 = (人工评分四项相加 ÷ 4)，请教师自行填写，可优先参考明确量化标准的机器评分，比如字数不足给5.0分，其次参考AI评分")

    doc.save(output_path)